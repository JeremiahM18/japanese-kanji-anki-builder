from __future__ import annotations

from datetime import datetime
from pathlib import Path
import json
import math
import re
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[2]
OUT_DIR = ROOT / "outputs" / "documents-whole-program-employer-packet"
FIG_DIR = OUT_DIR / "figures"
DOCX_PATH = OUT_DIR / "jkb_employer_reviewer_packet.docx"

BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
NAVY = RGBColor(0x0B, 0x25, 0x45)
MUTED = RGBColor(0x55, 0x65, 0x76)
INK = RGBColor(0x11, 0x18, 0x27)
GREEN = RGBColor(0x16, 0x65, 0x34)
AMBER = RGBColor(0x92, 0x5E, 0x00)
RED = RGBColor(0x9B, 0x1C, 0x1C)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
LIGHT_AMBER = "FFF7E6"
LIGHT_RED = "FDECEC"
LIGHT_GREEN = "EAF7EE"

DXA_PER_INCH = 1440
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def rgb_hex(color: RGBColor) -> str:
    return f"{color[0]:02X}{color[1]:02X}{color[2]:02X}"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, bottom=80, start=120, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        elem = tc_mar.find(qn(f"w:{side}"))
        if elem is None:
            elem = OxmlElement(f"w:{side}")
            tc_mar.append(elem)
        elem.set(qn("w:w"), str(value))
        elem.set(qn("w:type"), "dxa")


def set_paragraph_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_geometry(table, widths: list[int], indent_dxa: int = TABLE_INDENT_DXA) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx] / DXA_PER_INCH)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def keep_row_together(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def repeat_header_row(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:tblHeader")) is None:
        tbl_header = OxmlElement("w:tblHeader")
        tbl_header.set(qn("w:val"), "true")
        tr_pr.append(tbl_header)


def set_run_font(run, size: float | None = None, color: RGBColor | None = None, bold=None, italic=None, name="Calibri") -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_para_spacing(paragraph, before=0, after=6, line=1.10) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_run(paragraph, text: str, size=11, color=INK, bold=False, italic=False) -> None:
    run = paragraph.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold, italic=italic)


def add_body(doc: Document, text: str, after=6, color=INK) -> None:
    p = doc.add_paragraph()
    set_para_spacing(p, after=after)
    add_run(p, text, size=11, color=color)


def add_small(doc: Document, text: str, after=4, color=MUTED, italic=False) -> None:
    p = doc.add_paragraph()
    set_para_spacing(p, after=after)
    add_run(p, text, size=9.5, color=color, italic=italic)


def add_heading(doc: Document, text: str, level: int, page_break_before: bool = False) -> None:
    style = f"Heading {level}"
    p = doc.add_paragraph(style=style)
    p.paragraph_format.page_break_before = page_break_before
    p.text = text


def add_callout(doc: Document, label: str, text: str, fill: str = LIGHT_BLUE, color: RGBColor = NAVY) -> None:
    p = doc.add_paragraph()
    set_para_spacing(p, before=4, after=8, line=1.15)
    p.paragraph_format.left_indent = Pt(6)
    p.paragraph_format.right_indent = Pt(6)
    set_paragraph_shading(p, fill)
    add_run(p, f"{label}: ", size=10.5, color=color, bold=True)
    add_run(p, text, size=10.5, color=INK)


def add_picture_with_alt(doc: Document, path: Path, width, title: str, descr: str) -> None:
    shape = doc.add_picture(str(path), width=width)
    doc_pr = shape._inline.docPr
    doc_pr.set("title", title)
    doc_pr.set("descr", descr)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int], header_fill: str = LIGHT_GRAY, font_size=9.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    repeat_header_row(table.rows[0])
    keep_row_together(table.rows[0])
    for idx, header in enumerate(headers):
        cell = table.cell(0, idx)
        set_cell_shading(cell, header_fill)
        p = cell.paragraphs[0]
        set_para_spacing(p, after=0, line=1.05)
        add_run(p, header, size=font_size, color=NAVY, bold=True)
    for row_values in rows:
        row = table.add_row()
        keep_row_together(row)
        for idx, value in enumerate(row_values):
            cell = row.cells[idx]
            p = cell.paragraphs[0]
            set_para_spacing(p, after=0, line=1.05)
            add_run(p, value, size=font_size, color=INK)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def set_document_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.10


def set_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_para_spacing(header, after=0)
    add_run(header, "Japanese Kanji Builder - Employer/Reviewer Packet", size=9, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_para_spacing(footer, after=0)
    add_run(footer, "Generated evidence packet - not a release certificate", size=8.5, color=MUTED)


def wrap_text(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.ImageFont, width: int) -> list[str]:
    words = text.split()
    lines: list[str] = []
    current = ""
    for word in words:
        candidate = word if not current else f"{current} {word}"
        if draw.textbbox((0, 0), candidate, font=font)[2] <= width:
            current = candidate
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def font(size: int, bold=False):
    candidates = [
        "C:/Windows/Fonts/segoeuib.ttf" if bold else "C:/Windows/Fonts/segoeui.ttf",
        "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
    ]
    for candidate in candidates:
        try:
            return ImageFont.truetype(candidate, size)
        except OSError:
            continue
    return ImageFont.load_default()


def draw_box(draw, xy, title, body, fill, outline="#1F4D78") -> None:
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle(xy, radius=16, fill=fill, outline=outline, width=3)
    title_font = font(24, bold=True)
    body_font = font(19)
    draw.text((x1 + 18, y1 + 16), title, fill="#0B2545", font=title_font)
    y = y1 + 54
    for line in wrap_text(draw, body, body_font, x2 - x1 - 36):
        draw.text((x1 + 18, y), line, fill="#111827", font=body_font)
        y += 25


def draw_arrow(draw, start, end, color="#2E74B5") -> None:
    draw.line([start, end], fill=color, width=5)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    length = 16
    left = (end[0] - length * math.cos(angle - math.pi / 6), end[1] - length * math.sin(angle - math.pi / 6))
    right = (end[0] - length * math.cos(angle + math.pi / 6), end[1] - length * math.sin(angle + math.pi / 6))
    draw.polygon([end, left, right], fill=color)


def create_architecture_figures() -> tuple[Path, Path]:
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    fig1 = FIG_DIR / "architecture_pipeline.png"
    img = Image.new("RGB", (1800, 820), "white")
    draw = ImageDraw.Draw(img)
    title = "Governed Content Generation Pipeline"
    draw.text((60, 42), title, fill="#0B2545", font=font(36, bold=True))
    draw.text((60, 92), "Distribution output is Anki content; trust comes from contracts, gates, proof ledgers, and release boundaries.", fill="#374151", font=font(21))
    boxes = [
        ((70, 190, 350, 385), "Inputs", "JLPT contracts, curated data, media policy, source evidence, NLP support artifacts", "#E8EEF5"),
        ((430, 190, 710, 385), "Normalize", "Stable kanji identities and word|reading identities; level and product scope", "#F2F4F7"),
        ((790, 190, 1070, 385), "Generate", "Deterministic kanji and word rows; TSV and package-ready surfaces", "#EAF7EE"),
        ((1150, 190, 1430, 385), "Gate", "Gold regression, Sapphire structure, Platinum surface inspection", "#FFF7E6"),
        ((1510, 190, 1790, 385), "Prove", "Canonical JSONL Obsidian proof, reconciliation, scoped release locks", "#FDECEC"),
    ]
    for xy, title, body, fill in boxes:
        draw_box(draw, xy, title, body, fill)
    for x in [350, 710, 1070, 1430]:
        draw_arrow(draw, (x + 18, 288), (x + 70, 288))
    draw_box(draw, (430, 520, 810, 705), "Authority Boundary", "No lane borrows authority: source evidence, generated checks, review evidence, NLP support, media identity, and proof ledger certification stay separate.", "#F8FAFC")
    draw_box(draw, (990, 520, 1370, 705), "Release Boundary", "Deck/APKG readiness is mechanical support. Release trust still requires manual import, media, accessibility, listening, hosted checks, and blockers closed.", "#F8FAFC")
    draw_arrow(draw, (880, 612), (960, 612), "#64748B")
    img.save(fig1)

    fig2 = FIG_DIR / "lane_scope_map.png"
    img2 = Image.new("RGB", (1800, 980), "white")
    draw2 = ImageDraw.Draw(img2)
    draw2.text((60, 42), "Product Surfaces And Lane Scope", fill="#0B2545", font=font(36, bold=True))
    draw2.text((60, 92), "Whole-program status from live commands on 2026-06-11. Counts are generated denominators, not release claims.", fill="#374151", font=font(21))

    left_x = 90
    right_x = 940
    draw_box(draw2, (left_x, 180, left_x + 720, 375), "Core Kanji", "Generated: 2212/2212 across N5-N1. Locked Obsidian scope: 982/982 for N5-N2. N1 remains generated/Gold with 328/1230 Sapphire and Platinum.", "#E8EEF5")
    draw_box(draw2, (right_x, 180, right_x + 720, 375), "Word Decks", "Generated: 1470/1470 across N5-N1. Locked Obsidian scope: 987/987 for N5-N4. N3/N2/N1 word lanes remain forward backlog.", "#EAF7EE")
    ladder = [
        ("Silver", "Generated card surface exists"),
        ("Gold", "Regression-protected output"),
        ("Sapphire", "Structural certification"),
        ("Platinum", "Card-surface inspection"),
        ("Obsidian", "Explicit proof-ledger certification"),
    ]
    y = 485
    for idx, (lane, meaning) in enumerate(ladder):
        x = 330 + idx * 245
        draw_box(draw2, (x, y, x + 210, y + 132), lane, meaning, "#F8FAFC", "#2E74B5")
        if idx < len(ladder) - 1:
            draw_arrow(draw2, (x + 210, y + 66), (x + 245, y + 66), "#2E74B5")
    draw_box(draw2, (165, 725, 760, 895), "Support-Only Automation", "NLP governance passes and can assist review, but it does not certify cards, write tracked templates, or claim release readiness.", "#FFF7E6", "#925E00")
    draw_box(draw2, (1040, 725, 1635, 895), "Visible Backlog", "Expected coverage failures remain visible: N3/N2/N1 word forward lanes, plus N1 kanji Sapphire/Platinum backlog. Denominators are not shrunk.", "#FDECEC", "#9B1C1C")
    img2.save(fig2)
    return fig1, fig2


def read_doc_snippets() -> dict[str, str]:
    files = {
        "README": ROOT / "README.md",
        "Employer Overview": ROOT / "docs" / "employer-overview.md",
        "System Architecture": ROOT / "docs" / "system-architecture.md",
        "Release Lock": ROOT / "docs" / "releases" / "v0.2.0-scoped-obsidian-lock.md",
    }
    snippets = {}
    for key, path in files.items():
        snippets[key] = path.read_text(encoding="utf-8", errors="replace")[:2000]
    return snippets


def make_doc() -> None:
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    fig1, fig2 = create_architecture_figures()
    docs = read_doc_snippets()

    doc = Document()
    set_document_styles(doc)
    set_header_footer(doc)

    # Title block: standard business brief, memo masthead flavor.
    p = doc.add_paragraph()
    set_para_spacing(p, before=18, after=4)
    add_run(p, "EMPLOYER / REVIEWER PACKET", size=22, color=NAVY, bold=True)
    p = doc.add_paragraph()
    set_para_spacing(p, after=14)
    add_run(p, "Japanese Kanji Anki Builder", size=14, color=MUTED, bold=True)

    metadata = [
        ("Program identity", "Governed data pipeline / release-controlled content generation system"),
        ("Distribution artifact", "JLPT kanji and vocabulary study content for Anki"),
        ("Evidence mode", "Read-only document generation from live commands and tracked docs"),
        ("Generated", "2026-06-11 America/Chicago"),
        ("Repo commit", "a824d0bc data(words): expand n3 silver batches 35-41"),
        ("Scope", "Whole program: kanji and word decks, proof/governance, release-trust boundaries"),
    ]
    for label, value in metadata:
        p = doc.add_paragraph()
        set_para_spacing(p, after=2)
        add_run(p, f"{label}: ", size=10.5, color=INK, bold=True)
        add_run(p, value, size=10.5, color=INK)

    add_callout(
        doc,
        "Bottom line",
        "This repository is not merely a deck. It is a governed content data system that keeps generated surfaces, review lanes, proof ledgers, media identity, source evidence, NLP support, and release claims separate.",
        fill=LIGHT_BLUE,
    )

    add_heading(doc, "1. Sixty-Second Read", 1)
    add_body(
        doc,
        "A reviewer should understand the project as backend/data/product engineering: a controlled system that turns tracked contracts and curated local inputs into deterministic study artifacts, then protects trust claims through explicit gates and proof ledgers."
    )
    add_table(
        doc,
        ["Signal", "Verified state", "Why it matters"],
        [
            ["System type", "Governed data pipeline and release-controlled content generation system.", "Evaluates as software/data architecture, not a static flashcard file."],
            ["Whole-program scale", "2212 generated kanji rows and 1470 generated word rows across JLPT N5-N1.", "The system covers multiple product surfaces and keeps denominators visible."],
            ["Locked scope", "v0.2.0 freezes N5-N2 core kanji and N5-N4 words as scoped Obsidian-certified release content.", "Release claims are bounded instead of implied across unfinished work."],
            ["Proof storage", "Canonical JSONL ledger validates with 1969 events across 6 files.", "Certification evidence is tracked, queryable, and replayable."],
            ["Support boundary", "NLP governance passes, but it cannot certify cards or write tracked templates.", "Automation supports review without becoming unchecked authority."],
        ],
        [1700, 4000, 3660],
        font_size=9,
    )

    add_heading(doc, "2. Architecture", 1)
    add_body(doc, "The pipeline separates input authority, generated surfaces, review certification, proof storage, and release packaging. The diagrams below are derived from tracked architecture docs and live command posture.")
    add_picture_with_alt(
        doc,
        fig1,
        Inches(6.45),
        "Governed content generation pipeline",
        "Architecture diagram showing inputs, normalization, deterministic generation, review gates, canonical proof storage, authority boundaries, and release boundaries.",
    )
    add_small(doc, "Figure 1. Governed content generation pipeline. Source: docs/system-architecture.md plus live closeout/proof commands.", italic=True)
    add_picture_with_alt(
        doc,
        fig2,
        Inches(6.45),
        "Product surfaces and lane scope",
        "Diagram separating core kanji and word deck surfaces, Silver through Obsidian lanes, support-only automation, and visible unfinished backlog.",
    )
    add_small(doc, "Figure 2. Product surfaces and lane scope. Counts use live commands from this run; stale tracked docs still need count refresh where noted.", italic=True)

    add_heading(doc, "3. Current Product State", 1)
    add_table(
        doc,
        ["Surface", "Generated", "Gold", "Sapphire", "Platinum", "Obsidian / proof boundary"],
        [
            ["Kanji N5", "80/80", "80/80", "80/80", "80/80", "Locked scope; Obsidian certified."],
            ["Kanji N4", "212/212", "212/212", "212/212", "212/212", "Locked scope; Obsidian certified."],
            ["Kanji N3", "341/341", "341/341", "341/341", "341/341", "Locked scope; Obsidian certified."],
            ["Kanji N2", "349/349", "349/349", "349/349", "349/349", "Locked scope; Obsidian certified."],
            ["Kanji N1", "1230/1230", "1230/1230", "328/1230", "328/1230", "No trusted Obsidian proof counted; 902 rows need current Sapphire and Platinum first."],
            ["Word N5", "287/287", "287/287", "287/287", "287/287", "Locked scope; Obsidian certified."],
            ["Word N4", "700/700", "700/700", "700/700", "700/700", "Locked scope; Obsidian certified."],
            ["Word N3", "429/429", "8/429", "8/429", "8/429", "Forward backlog; Obsidian proof not recorded."],
            ["Word N2", "28/28", "0/28", "0/28", "0/28", "Forward backlog; Obsidian not started."],
            ["Word N1", "26/26", "0/26", "0/26", "0/26", "Forward backlog; Obsidian not started."],
        ],
        [1280, 1050, 1050, 1150, 1150, 3680],
        font_size=8.5,
    )

    add_heading(doc, "4. Proof And Governance State", 1)
    add_table(
        doc,
        ["Check", "Result", "Evidence"],
        [
            ["Proof ledger validation", "PASS", "1969 proof events across 6 JSONL files."],
            ["Kanji proof reconciliation N5-N2", "PASS", "982 ledger proofs bound; 0 mismatches."],
            ["Word proof reconciliation N5-N4", "PASS", "987 ledger proofs bound; 0 mismatches."],
            ["Kanji Obsidian certification N5-N2", "PASS", "982/982 Obsidian certified; 0 needs Obsidian; 0 blocked/failing."],
            ["Word Obsidian certification N5-N4", "PASS", "987/987 Obsidian certified; 0 needs Obsidian; 0 blocked/failing."],
            ["NLP governance gate", "PASS", "8 checks pass; certifies cards: no; writes tracked templates: no."],
            ["Performance/memory matrix", "PASS", "7 lanes pass as governance map; not a release-readiness claim."],
            ["Release trust", "FAIL-CLOSED", "Open release blockers: SEC-P0-004, PROD-REL-001, SEC-REQ-007."],
        ],
        [2500, 1500, 5360],
        font_size=9,
    )

    add_callout(
        doc,
        "Important limitation",
        "The proof ledger and Obsidian status commands certify scoped content proof only. They do not prove APKG import, mobile behavior, accessibility, listening quality, hosted CI status, source-evidence depth completion, or broader release readiness.",
        fill=LIGHT_AMBER,
        color=AMBER,
    )

    add_heading(doc, "5. Word Reading And Expansion Posture", 1, page_break_before=True)
    add_table(
        doc,
        ["Level", "Inventory", "Built", "Readiness", "Reading coverage", "Notes"],
        [
            ["N5", "287", "287/287", "ready; deferred variants", "67.7% (233/344)", "Active triage cleared; exact command label preserved in appendix."],
            ["N4", "700", "700/700", "ready; deferred variants", "76.7% (579/755)", "Active triage cleared; exact command label preserved in appendix."],
            ["N3", "429", "429/429", "incomplete", "34.9% (413/1184)", "Silver expansion exists; forward review lanes are 8/429."],
            ["N2", "28", "28/28", "incomplete", "5.4% (56/1045)", "Forward review lanes not started."],
            ["N1", "26", "26/26", "incomplete", "1.3% (41/3274)", "Forward review lanes not started."],
        ],
        [800, 1100, 950, 1800, 1650, 3060],
        font_size=8.5,
    )

    add_heading(doc, "6. Explicit Limitations", 1)
    limitations = [
        ("Tracked-doc count drift", "docs/employer-overview.md and docs/system-architecture.md still say 1400 word rows. Live commands now show 1470. Update tracked docs before sending them as the sole external evidence."),
        ("Generated artifact dirtiness", "git status is not clean because generated outputs/ artifacts exist. Tracked source excluding outputs remains clean."),
        ("Release readiness", "security:release-trust fails closed on two high/critical release-blocker risks and one unimplemented release-blocker requirement."),
        ("Manual QA", "APKG import, mobile, accessibility, listening, and managed-media review are not proven by these commands."),
        ("NLP", "NLP artifacts are support-only. Passing NLP governance does not certify a card or mutate tracked templates."),
        ("Unfinished lanes", "Expected coverage failures remain visible for N1 kanji Sapphire/Platinum and N3/N2/N1 word Gold/Sapphire/Platinum."),
    ]
    add_table(doc, ["Limitation", "Verified handling"], [[a, b] for a, b in limitations], [2100, 7260], font_size=9)

    add_heading(doc, "7. Reviewer Route", 1)
    add_body(doc, "Use this packet as an orientation layer, then verify with the commands in the appendix. The strongest tracked human-facing docs are README.md, docs/employer-overview.md, docs/system-architecture.md, docs/releases/v0.2.0-scoped-obsidian-lock.md, docs/product-exit-criteria.md, docs/verification.md, and docs/release-process.md.")
    add_body(doc, "For a hiring reviewer, the most important engineering signal is not the existence of cards. It is the system's refusal to collapse generated output, structural checks, learner-surface inspection, proof-ledger certification, and release readiness into one vague claim.")

    add_heading(doc, "Appendix A. Live Command Evidence", 1, page_break_before=True)
    command_rows = [
        ["git status --short --branch", "PASS", "main tracking origin/main; generated outputs/ present."],
        ["git status --short --branch -- . :(exclude)outputs", "PASS", "Tracked source clean when generated outputs are excluded."],
        ["git log -1 --oneline --decorate", "PASS", "a824d0bc data(words): expand n3 silver batches 35-41."],
        ["git ls-remote --heads origin", "PASS", "Remote heads show main at a824d0bc5ee8f405cf3aacf2e00ac111251b0cf7."],
        ["npm run deck:closeout -- --levels=5,4,3,2,1", "PASS / expected gaps", "Reports whole-program lane counts; expected coverage failures remain visible."],
        ["npm run data:obsidian:proof:validate", "PASS", "1969 proof events, 6 files."],
        ["npm run data:obsidian:proof:reconcile -- --levels=5,4,3,2", "PASS", "982 kanji proofs bound, 0 mismatches."],
        ["npm run data:obsidian:proof:reconcile -- --deck-kind=word --levels=5,4", "PASS", "987 word proofs bound, 0 mismatches."],
        ["npm run deck:kanji:obsidian:certify-status -- --levels=5,4,3,2", "PASS", "982/982 certified for locked kanji scope."],
        ["npm run deck:words:obsidian:certify-status -- --levels=5,4", "PASS", "987/987 certified for locked word scope."],
        ["npm run nlp:governance-gate", "PASS", "8 checks pass; support-only boundaries explicit."],
        ["npm run perf:memory:matrix", "PASS", "7 governance lanes pass; benchmark authority remains separate."],
        ["npm run security:release-trust", "EXPECTED FAIL-CLOSED", "SEC-P0-004, PROD-REL-001, and SEC-REQ-007 remain blockers."],
        ["npm run deck:words:completion:n5", "PASS", "N5 ready_with_deferred_variants; reading coverage 67.7%."],
        ["npm run deck:words:completion:n4", "PASS", "N4 ready_with_deferred_variants; reading coverage 76.7%."],
        ["npm run deck:words:completion:n3", "PASS / incomplete", "N3 generated 429/429; reading coverage 34.9%; forward lanes 8/429."],
        ["npm run deck:words:completion:n2", "PASS / incomplete", "N2 generated 28/28; reading coverage 5.4%; forward lanes not started."],
        ["npm run deck:words:completion:n1", "PASS / incomplete", "N1 generated 26/26; reading coverage 1.3%; forward lanes not started."],
    ]
    add_table(doc, ["Command", "Classification", "Packet use"], command_rows, [3500, 1600, 4260], font_size=8)

    add_heading(doc, "Appendix B. Tracked Document Sources Read", 1)
    source_rows = [
        ["README.md", "Project identity, authority boundary, review tiers, current baseline, source-of-truth map."],
        ["docs/employer-overview.md", "Employer-facing framing; count drift noted against live commands."],
        ["docs/system-architecture.md", "Architecture maps and lane authority model; count drift noted against live commands."],
        ["docs/releases/v0.2.0-scoped-obsidian-lock.md", "Frozen release scope and release-ready caveats."],
        ["package.json", "Command inventory and program description."],
        ["outputs/spreadsheets-whole-program-governance", "Prior workbook/export contract and visual dashboard previews."],
    ]
    add_table(doc, ["Source", "How used"], source_rows, [2700, 6660], font_size=9)

    add_small(doc, "Boundary statement: This packet is an evidence-backed orientation artifact. It does not certify new cards, append proof, change lane counts, update tracked docs, approve release readiness, or replace rerunning the commands above at decision time.", italic=True)

    # Persist an evidence summary next to the DOCX for auditability.
    evidence = {
        "generatedAt": datetime.now().isoformat(timespec="seconds"),
        "docxPath": str(DOCX_PATH),
        "liveCounts": {
            "coreKanjiGenerated": 2212,
            "wordGenerated": 1470,
            "lockedKanjiObsidian": 982,
            "lockedWordObsidian": 987,
            "proofEvents": 1969,
        },
        "knownDocDrift": ["docs/employer-overview.md word denominator 1400", "docs/system-architecture.md word denominator 1400"],
        "figures": [str(fig1), str(fig2)],
    }
    (OUT_DIR / "packet_evidence_summary.json").write_text(json.dumps(evidence, indent=2), encoding="utf-8")
    doc.save(DOCX_PATH)


if __name__ == "__main__":
    make_doc()
    print(DOCX_PATH)
